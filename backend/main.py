from fastapi import FastAPI, UploadFile, File, Response
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse, JSONResponse, FileResponse
from pydantic import BaseModel
from typing import Optional
import httpx
import chromadb
from chromadb.utils import embedding_functions
import sqlite3
import uuid
import re
import base64
import io
from datetime import datetime, timezone
import math
from ddgs import DDGS
from pypdf import PdfReader
from sheets import get_menu_context, get_recipes, cross_reference_ingredients
from shopify_api import create_meal_products as _create_meal_products
from elevenlabs.client import ElevenLabs
import os
import whisper
import tempfile
import json
from pathlib import Path
from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import threading
from dotenv import load_dotenv
load_dotenv()

# ── App setup ────────────────────────────────────────────────────────────────
app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=[
        "http://localhost:5173",
        "http://192.168.1.112:5173",
        "https://lucchese.app",
        "https://www.lucchese.app",
    ],
    allow_methods=["*"],
    allow_headers=["*"],
)

DOCS_DIR = Path("./generated_docs")
DOCS_DIR.mkdir(exist_ok=True)

_doc_store: dict[str, str] = {}
_doc_lock = threading.Lock()

# ── ChromaDB setup ───────────────────────────────────────────────────────────
CHROMA_PATH   = "./chroma_db"
EMBED_MODEL   = "all-MiniLM-L6-v2"
CHUNK_SIZE    = 800
CHUNK_OVERLAP = 100

ef = embedding_functions.SentenceTransformerEmbeddingFunction(model_name=EMBED_MODEL)
chroma = chromadb.PersistentClient(path=CHROMA_PATH)

col_knowledge = chroma.get_or_create_collection("knowledge", embedding_function=ef)
col_facts     = chroma.get_or_create_collection("facts",     embedding_function=ef)
col_style     = chroma.get_or_create_collection("style",     embedding_function=ef)
col_documents = chroma.get_or_create_collection("documents", embedding_function=ef)
col_summaries = chroma.get_or_create_collection("summaries", embedding_function=ef)


# ── Ollama config ────────────────────────────────────────────────────────────
OLLAMA_URL = os.getenv("OLLAMA_BASE_URL", "http://localhost:11434") + "/api/chat"
MODEL_FAST = os.getenv("MODEL_FAST", "gemma2:27b")
MODEL_DEEP = os.getenv("MODEL_DEEP", "qwen2.5:32b")
ANTHROPIC_API_KEY = os.getenv("ANTHROPIC_API_KEY", "")
CHAT_PROVIDER     = os.getenv("CHAT_PROVIDER", "ollama")
print(f"DEBUG — Provider: {CHAT_PROVIDER}, Key set: {bool(ANTHROPIC_API_KEY)}")  # remove after testing

ELEVENLABS_API_KEY = os.getenv("ELEVENLABS_API_KEY")
ELEVENLABS_VOICE_ID = os.getenv("ELEVENLABS_VOICE_ID")
el_client = ElevenLabs(api_key=ELEVENLABS_API_KEY)

whisper_model = whisper.load_model("tiny")


# ── SQLite setup ─────────────────────────────────────────────────────────────
DB_PATH = "./conversations.db"

def init_db():
    con = sqlite3.connect(DB_PATH)
    cur = con.cursor()
    cur.execute("""
        CREATE TABLE IF NOT EXISTS conversations (
            id          TEXT PRIMARY KEY,
            created_at  TEXT NOT NULL,
            updated_at  TEXT NOT NULL,
            title       TEXT
        )
    """)
    cur.execute("""
        CREATE TABLE IF NOT EXISTS messages (
            id              TEXT PRIMARY KEY,
            conversation_id TEXT NOT NULL,
            role            TEXT NOT NULL,
            content         TEXT NOT NULL,
            created_at      TEXT NOT NULL,
            FOREIGN KEY (conversation_id) REFERENCES conversations(id)
        )
    """)
    cur.execute("""
        CREATE TABLE IF NOT EXISTS documents (
            id          TEXT PRIMARY KEY,
            filename    TEXT NOT NULL,
            file_type   TEXT NOT NULL,
            chunk_count INTEGER NOT NULL,
            created_at  TEXT NOT NULL
        )
    """)
    cur.execute("CREATE INDEX IF NOT EXISTS idx_messages_conv_id ON messages(conversation_id)")
    cur.execute("""
        CREATE TABLE IF NOT EXISTS roleplay_sessions (
            conversation_id TEXT PRIMARY KEY,
            exchanges       INTEGER NOT NULL DEFAULT 0,
            created_at      TEXT NOT NULL
        )
    """)
    con.commit()
    con.close()

init_db()

def _add_inline_formatting(para, text: str):
    """Handle **bold** and *italic* inline markdown within a paragraph."""
    pattern = re.compile(r"(\*\*.*?\*\*|\*.*?\*)")
    parts   = pattern.split(text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run      = para.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run        = para.add_run(part[1:-1])
            run.italic = True
        else:
            para.add_run(part)


def _looks_like_section_header(line: str) -> bool:
    """
    Detect plain-text section headers like "Property Basics:" or "Exit Strategy:"
    when the model doesn't use markdown heading syntax.
    Rules: ends with colon, 1-5 words, no sentence punctuation, starts with capital.
    """
    s = line.strip()
    if not s.endswith(":"):
        return False
    if len(s) > 50:
        return False
    if not s[0].isupper():
        return False
    # 1-5 words only — section names, not sentences
    word_count = len(s.rstrip(":").split())
    if word_count > 5:
        return False
    # No sentence punctuation inside (rules out "Let's break this down:")
    if re.search(r"[,\.!?']", s.rstrip(":")):
        return False
    return True


def markdown_to_docx(content: str, title: str) -> str:
    """
    Convert markdown-ish text to a .docx file.
    Handles: # ## ### headings, plain "Section:" headers, - bullets,
    1. numbered lists, **bold**, *italic*, --- rules.
    Returns the filepath of the saved document.
    """
    doc = DocxDocument()

    # ── Page margins ──────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.2)
        section.right_margin  = Inches(1.2)

    # ── Title ─────────────────────────────────────────────────────────────────
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title_para.add_run(title.replace("_", " ").title())
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

    date_para = doc.add_paragraph()
    date_run  = date_para.add_run(
        datetime.now(timezone.utc).strftime("Generated %d %B %Y")
    )
    date_run.font.size = Pt(9)
    date_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    doc.add_paragraph()  # spacer

    # ── Body ──────────────────────────────────────────────────────────────────
    for line in content.split("\n"):
        s = line.strip()

        if not s:
            doc.add_paragraph()
            continue

        # H1: # Heading
        if s.startswith("# "):
            p   = doc.add_paragraph()
            run = p.add_run(s[2:].strip())
            run.bold = True
            run.font.size = Pt(15)
            run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after  = Pt(4)

        # H2: ## Heading
        elif s.startswith("## "):
            p   = doc.add_paragraph()
            run = p.add_run(s[3:].strip())
            run.bold = True
            run.font.size = Pt(13)
            run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after  = Pt(3)

        # H3: ### Heading
        elif s.startswith("### "):
            p   = doc.add_paragraph()
            run = p.add_run(s[4:].strip())
            run.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x55)
            p.paragraph_format.space_before = Pt(8)

        # Plain-text section header fallback: "Section Name:"
        elif _looks_like_section_header(s):
            p   = doc.add_paragraph()
            run = p.add_run(s.rstrip(":"))
            run.bold = True
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after  = Pt(3)

        # Bullet: - item or * item
        elif s.startswith(("- ", "* ")):
            p = doc.add_paragraph(style="List Bullet")
            _add_inline_formatting(p, s[2:].strip())

        # Numbered: 1. item
        elif re.match(r"^\d+\.\s", s):
            p = doc.add_paragraph(style="List Number")
            _add_inline_formatting(p, re.sub(r"^\d+\.\s*", "", s))

        # Horizontal rule
        elif s in ("---", "***", "___"):
            p   = doc.add_paragraph()
            run = p.add_run("─" * 60)
            run.font.color.rgb = RGBColor(0xcc, 0xcc, 0xcc)
            run.font.size = Pt(8)

        # Normal paragraph
        else:
            p = doc.add_paragraph()
            _add_inline_formatting(p, s)
            p.paragraph_format.space_after = Pt(4)

    # ── Save ──────────────────────────────────────────────────────────────────
    safe_name = re.sub(r"[^\w\-]", "_", title)[:60]
    filename  = f"{safe_name}_{uuid.uuid4().hex[:6]}.docx"
    filepath  = DOCS_DIR / filename
    doc.save(str(filepath))
    return str(filepath)



# ── DB helpers ────────────────────────────────────────────────────────────────
def get_con():
    con = sqlite3.connect(DB_PATH)
    con.row_factory = sqlite3.Row
    return con

def save_message(conversation_id: str, role: str, content: str):
    now = datetime.now(timezone.utc).isoformat()
    con = get_con()
    cur = con.cursor()
    try:
        cur.execute("""
            INSERT INTO conversations (id, created_at, updated_at, title)
            VALUES (?, ?, ?, NULL)
            ON CONFLICT(id) DO UPDATE SET updated_at = excluded.updated_at
        """, (conversation_id, now, now))

        cur.execute(
            "INSERT INTO messages (id, conversation_id, role, content, created_at) VALUES (?, ?, ?, ?, ?)",
            (str(uuid.uuid4()), conversation_id, role, content, now)
        )

        cur.execute("""
            UPDATE conversations SET title = (
                SELECT substr(content, 1, 60) || CASE WHEN length(content) > 60 THEN '...' ELSE '' END
                FROM messages WHERE conversation_id = ? AND role = 'user'
                ORDER BY created_at LIMIT 1
            ) WHERE id = ?
        """, (conversation_id, conversation_id))

        con.commit()
    except Exception as e:
        print(f"save_message error: {e}")
        con.rollback()
    finally:
        con.close()

# ── Text chunking ─────────────────────────────────────────────────────────────
def chunk_text(text: str) -> list[str]:
    # Split on sentence boundaries first
    sentences = re.split(r'(?<=[.!?])\s+', text.strip())
    chunks = []
    current = ""

    for sentence in sentences:
        if len(current) + len(sentence) + 1 <= CHUNK_SIZE:
            current = (current + " " + sentence).strip()
        else:
            if current:
                chunks.append(current)
            # If a single sentence exceeds chunk size, split it hard
            if len(sentence) > CHUNK_SIZE:
                for i in range(0, len(sentence), CHUNK_SIZE - CHUNK_OVERLAP):
                    part = sentence[i:i + CHUNK_SIZE]
                    if part.strip():
                        chunks.append(part.strip())
            else:
                current = sentence

    if current:
        chunks.append(current)

    return [c for c in chunks if c.strip()]

def classify_memory(text: str) -> str:
    t = text.lower()

    # ── PTPreps Business ──────────────────────────────────────────────────────
    if any(w in t for w in ['ptpreps', 'pt preps', 'meal prep business', 'subscription', 'shopify', 'appstle', 'allergen label', 'delivery checker', 'bundle', 'cook plan', 'fulfilment', 'dispatch', 'packaging', 'refund', 'allergen', 'haccp', 'food safety', 'food cooling', 'allergen listing']):
        return 'ptpreps'

    # ── Personal Training & Coaching ──────────────────────────────────────────
    if any(w in t for w in ['pt client', 'personal training', 'coaching client', 'programme card', 'check in', 'l3', 'level 3', 'practical assessment', 'smart goal', 'pt course', 'puregym assessment', 'phil feedback']):
        return 'coaching'

    # ── Fitness & Training ────────────────────────────────────────────────────
    if any(w in t for w in ['gym', 'workout', 'training', 'bodybuilding', 'lift', 'squat', 'deadlift', 'bench', 'cardio', 'mma', 'bjj', 'boxing', 'sparring', 'martial', 'cut', 'bulk', 'physique', 'looksmaxxing', 'mind muscle', 'shin splint', 'grip strength', 'leg day', 'resistance band', 'smelling salt']):
        return 'fitness'

    # ── Food & Recipes ────────────────────────────────────────────────────────
    if any(w in t for w in ['recipe', 'macro', 'protein', 'kcal', 'calorie', 'ingredient', 'menu', 'portion', 'bulking recipe', 'standard recipe', 'meal plan', 'nutrition', 'food diary', 'raw weight', 'cook weight']):
        return 'food'

    if any(w in t for w in ['cook', 'bake', 'roast', 'fry', 'sauté', 'braise', 'cuisine', 'kitchen', 'chef', 'jerk chicken', 'beef bourguignon', 'venison', 'sourdough', 'tallow', 'gastro']):
        return 'cooking'

    # ── Health & Body ─────────────────────────────────────────────────────────
    if any(w in t for w in ['health', 'cholesterol', 'sleep', 'injury', 'recovery', 'blood test', 'seizure', 'epilepsy', 'pain', 'doctor', 'nhs', 'medication', 'levetiracetam', 'pip', 'fit note', 'universal credit', 'deviated septum', 'jaw', 'chest spasm', 'vape', 'vaping', 'nicotine', 'cpap', 'sleep apnea', 'uars', 'daylight saving', 'circadian']):
        return 'health'

    # ── Compounds, PEDs & Peptides ────────────────────────────────────────────
    if any(w in t for w in ['trt', 'testosterone', 'test e', 'test enanthate', 'deca', 'eq ', 'equipoise', 'hcg', 'aromatase', 'estrogen', 'steroid', 'sarm', 'peptide', 'bpc', 'bpc-157', 'tb-500', 'glute injection', 'methylene blue', 'epo', 'armodafinil', 'modafinil']):
        return 'compounds'

    if any(w in t for w in ['supplement', 'creatine', 'vitamin', 'omega', 'magnesium', 'zinc', 'pre workout', 'whey', 'collagen', 'semax', 'selank', 'nootropic', 'paul stamets', 'microdosing', 'lsd', 'mushroom stack', 'caffeine', 'l-theanine', 'nlp', 'cognitive']):
        return 'supplements_nootropics'

    # ── Psychedelics & Altered States ─────────────────────────────────────────
    if any(w in t for w in ['psilocybin', 'mushroom', 'shroom', 'mycelium', 'grow kit', 'fungi', 'dmt', 'ayahuasca', 'sensory deprivation', 'float tank', 'cannabis', 'weed', 'cocaine', 'drug', 'addiction', 'withdrawal', 'perception', 'altered state']):
        return 'psychedelics'

    # ── Property & Real Estate ────────────────────────────────────────────────
    if any(w in t for w in ['property', 'mortgage', 'equity', 'tenant', 'hmo', 'landlord', 'brrr', 'r2r', 'yield', 'bridging', 'remortgage', 'estate agent', 'conveyancer', 'sdlt', 'rent to rent', 'deal analysis', 'analyse deal']):
        return 'property'

    # ── Finance & Money ───────────────────────────────────────────────────────
    if any(w in t for w in ['invest', 'stock', 'share', 'crypto', 'bitcoin', 'portfolio', 'dividend', 'index fund', 'isa', 'pension', 'compound interest', 'asset', 'holding company', 'family investment', 'financial literacy', 'economy', 'wealth']):
        return 'finance'

    if any(w in t for w in ['budget', 'spend', 'afford', 'debt', 'loan', 'cash flow', 'salary', 'invoice', 'tax', 'hmrc', 'self assessment', 'ltd company', 'sole trader', 'revolut', 'payroll', 'director pay', 'universal credit']):
        return 'money'

    # ── Tech & Building ───────────────────────────────────────────────────────
    if any(w in t for w in ['chroma', 'ollama', 'embeddings', 'rag', 'vector', 'building lucchese', 'lucchese feature', 'lucchese bug', 'lucchese memory', 'lucchese db', 'ai dashboard', 'ingest']):
        return 'lucchese'

    if any(w in t for w in ['code', 'python', 'javascript', 'react', 'fastapi', 'api', 'backend', 'frontend', 'debug', 'server', 'deploy', 'vite', 'database', 'sqlite', 'github', 'npm', 'venv', 'flask', 'ngrok', 'uvicorn']):
        return 'tech'

    if any(w in t for w in ['hack', 'pentest', 'security', 'xss', 'recon', 'ctf', 'exploit', 'cyber', 'vulnvault', 'osint', 'phishing', 'bug bounty', 'threat intel', 'kali']):
        return 'security'

    if any(w in t for w in ['ai model', 'gpt', 'claude', 'llm', 'prompt', 'machine learning', 'automation', 'chatgpt', 'grok', 'hallucination']):
        return 'ai'

    if any(w in t for w in ['pc build', 'gpu', 'cpu', 'psu', 'nzxt', 'radiator', 'liquid cooler', 'monitor', 'arduino', 'elegoo', 'mini pc', 'usb', 'coaxial', 'daisy chain']):
        return 'hardware'

    # ── Content & Creative ────────────────────────────────────────────────────
    if any(w in t for w in ['reel', 'youtube', 'tiktok', 'instagram', 'video edit', 'short form', 'thumbnail', 'caption', 'viral', 'content creator', '60 second']):
        return 'content'

    if any(w in t for w in ['music', 'song', 'track', 'dj', 'vinyl', 'beat', 'produce', 'guitar', 'drum', 'playlist', 'album', 'artist', 'songwriting', 'mandolin', 'pas rudiment', 'warner classics']):
        return 'music'

    if any(w in t for w in ['film', 'movie', 'series', 'netflix', 'documentary', 'episode', 'rdr2', 'minecraft', 'gaming', 'console', 'playstation', 'xbox', 'steam', 'gta', 'bg3']):
        return 'media_gaming'

    # ── Psychology & Mind ─────────────────────────────────────────────────────
    if any(w in t for w in ['psychology', 'psychopathy', 'narcissism', 'trauma', 'attachment', 'ego', 'shadow', 'identity', 'subconscious', 'behaviour', 'habit', 'addiction cycle', 'breaking the cycle', 'family enmeshment', 'boundaries', 'parenting yourself', 'inner child']):
        return 'psychology'

    if any(w in t for w in ['mindset', 'discipline', 'motivation', 'self improvement', 'stoic', 'focus', 'productivity', 'confidence', 'resilience', 'procrastination', 'resistance', 'momentum', 'consistency', 'insanity vs consistency']):
        return 'mindset'

    if any(w in t for w in ['neuroscience', 'neuroplasticity', 'dopamine', 'serotonin', 'neuromelanin', 'nervous system', 'brain', 'neurology', 'cognitive', 'consciousness', 'layers of consciousness', 'mind muscle']):
        return 'neuroscience'

    if any(w in t for w in ['nlp', 'attraction', 'masculine', 'feminine', 'dating', 'relationship', 'social dynamics', 'gym talk', 'gym convo', 'social tension', 'nice guys', 'manosphere']):
        return 'social_dynamics'

    # ── Family & Personal ─────────────────────────────────────────────────────
    if any(w in t for w in ['mum', 'dad', 'brother', 'sister', 'aunt', 'uncle', 'family', 'parent', 'grandma', 'grandad', 'nan', 'siblings', 'family conflict', 'family chaos', 'family pressure']):
        return 'family'

    if any(w in t for w in ['friend', 'mate', 'lad', 'social', 'night out', 'energy audit', 'friendships', 'protecting energy']):
        return 'relationships'

    if any(w in t for w in ['travel', 'holiday', 'flight', 'hotel', 'abroad', 'country', 'city', 'trip', 'amazon adventure', 'fishing']):
        return 'travel'

    # ── Spirituality & Esoteric ───────────────────────────────────────────────
    if any(w in t for w in ['consciousness', 'simulation', 'spiritual', 'universe', 'frequency', 'ancient', 'dimension', 'mandela effect', 'pyramid', 'parallel universe', 'quantum', 'double slit', 'superposition', 'dmt burning bush', 'heaven', 'hell', 'ascension', 'astral']):
        return 'spirituality'

    if any(w in t for w in ['conspiracy', 'deep state', 'adrenochrome', 'epstein', 'diddy', 'cia', 'isis cia', 'hearst hemp', 'flat earth', 'bob lazar', 'greer', 'ufo', 'alien', 'buga sphere', 'giants', 'dinosaur hoax', 'soros', 'sumitomo']):
        return 'conspiracy'

    if any(w in t for w in ['god', 'faith', 'religion', 'church', 'belief', 'bible', 'jesus', 'muslim', 'christian', 'jewish', 'quran', 'fallen angel', 'esau', 'ptah', 'egypt', 'freemason', 'satanism', 'antichrist', 'black magic', 'immortality key', 'psychedelic religion']):
        return 'religion_esoteric'

    if any(w in t for w in ['philosophy', 'meaning', 'purpose', 'existence', 'free will', 'morality', 'ethics', 'truth', 'simulation theory', 'devil hypothetical', 'your prophecy']):
        return 'philosophy'

    # ── World & Society ───────────────────────────────────────────────────────
    if any(w in t for w in ['politic', 'government', 'vote', 'election', 'law', 'immigration', 'war', 'policy', 'tommy robinson', 'israel', 'palestine', 'zionism', 'anti-semitism', 'green party', 'getting into politics', 'uk demographics']):
        return 'politics'

    if any(w in t for w in ['economy', 'inflation', 'central bank', 'money printing', 'nationalise', 'capitalism', 'decentralised', 'digital currency', 'global tension', 'world leadership', 'cultural extreme']):
        return 'economics'

    if any(w in t for w in ['history', 'civilisation', 'empire', 'ancient egypt', 'pyramids', 'solzhenitsyn', 'george washington', 'freddie mercury', 'elon musk', 'dan pena', 'paul allen', 'theo von', 'louis theroux', 'bob marley']):
        return 'history_culture'

    if any(w in t for w in ['science', 'physics', 'biology', 'chemistry', 'comet', 'ice age', 'asteroid', 'light pollution', 'sky', 'quantum', 'seed oil', 'cholesterol debate']):
        return 'science'

    if any(w in t for w in ['autism', 'schizophrenia', 'epilepsy awareness', 'mental health', 'self loathing', 'anger', 'overthinking', 'breaking free', 'stuck', 'chaos', 'life direction', 'reset', 'getting back on track']):
        return 'mental_health'

    # ── Career & Study ────────────────────────────────────────────────────────
    if any(w in t for w in ['job', 'recruit', 'hire', 'salary', 'career', 'interview', 'cv', 'comptia', 'it cert', 'microsoft 365', 'it support', 'it role']):
        return 'career'

    return 'general'

async def ingest_exchange(conversation_id: str, user_msg: str, assistant_reply: str):
    now = datetime.now(timezone.utc).isoformat()
    exchange_text = f"[You]: {user_msg}\n\n[Lucchese]: {assistant_reply}"

    for i, chunk in enumerate(chunk_text(exchange_text)):
        if not is_duplicate_memory(chunk, col_knowledge):
            col_knowledge.upsert(
                ids       = [f"lucchese_{conversation_id}_ex_{i}"],
                documents = [chunk],
                metadatas = [{
                    "source":     "lucchese",
                    "conv_id":    conversation_id,
                    "chunk_idx":  str(i),
                    "type":       "knowledge",
                    "category":   classify_memory(chunk),
                    "created_at": now,
                }]
            )

    if len(user_msg.strip()) > 20:
        if not is_duplicate_memory(user_msg.strip(), col_facts):
            col_facts.upsert(
                ids       = [f"lucchese_fact_{conversation_id}_{uuid.uuid4().hex[:8]}"],
                documents = [user_msg.strip()],
                metadatas = [{
                    "source":     "lucchese",
                    "conv_id":    conversation_id,
                    "type":       "fact",
                    "category":   classify_memory(user_msg.strip()),
                    "created_at": now,
                }]
            )

    if len(user_msg.strip()) > 80:
        if not is_duplicate_memory(user_msg.strip(), col_style):
            col_style.upsert(
                ids       = [f"lucchese_style_{conversation_id}_{uuid.uuid4().hex[:8]}"],
                documents = [user_msg.strip()],
                metadatas = [{
                    "source":     "lucchese",
                    "conv_id":    conversation_id,
                    "type":       "style",
                    "category":   classify_memory(user_msg.strip()),
                    "created_at": now,
                }]
            )

# ── Memory quality filter ─────────────────────────────────────────────────────
def should_ingest(user_msg: str, assistant_reply: str) -> bool:
    msg = user_msg.lower().strip()
    
    # Too short to be meaningful
    if len(msg) < 30:
        return False
    
    # Assistant was uncertain — not worth storing
    uncertainty = ["i don't know", "i'm not sure", "i can't find", "i don't have"]
    if any(p in assistant_reply.lower() for p in uncertainty):
        return False
    
    # Strong signal words — definitely store
    signals = [
        "my ", "i am", "i'm", "we ", "i have", "i've", "i do", "i don't",
        "alex", "ptpreps", "prefer", "always", "never", "usually", "hate",
        "love", "want", "need", "goal", "plan", "think", "believe", "feel"
    ]
    if any(w in msg for w in signals):
        return True
    
    return False
    
    


# ── Memory deduplication ──────────────────────────────────────────────────────
SIMILARITY_THRESHOLD = 0.35  # lower = stricter dedup (ChromaDB uses L2 distance)

def is_duplicate_memory(text: str, collection, n=1) -> bool:
    """Returns True if a very similar chunk already exists in the collection."""
    try:
        results = collection.query(query_texts=[text], n_results=n, include=["distances"])
        if results["distances"] and results["distances"][0]:
            closest_distance = results["distances"][0][0]
            return closest_distance < SIMILARITY_THRESHOLD
    except Exception:
        pass
    return False

# ── Explicit memory commands ──────────────────────────────────────────────────
REMEMBER_PATTERNS = [
    r'\bremember that\b(.+)',
    r'\bremember:\b(.+)',
    r'\bsave this\b(.+)',
    r'\bstore this\b(.+)',
    r'\bnote that\b(.+)',
]

FORGET_PATTERNS = [
    r'\bforget that\b(.+)',
    r'\bforget\b (.+)',
    r'\bdelete that\b',
    r'\bunremember\b(.+)',
    r'\bremove that\b',
]

def detect_memory_command(message: str):
    """Returns ('remember', text), ('forget', text), or (None, None)"""
    msg = message.lower().strip()
    
    for pattern in REMEMBER_PATTERNS:
        match = re.search(pattern, msg)
        if match:
            return ("remember", match.group(1).strip())
    
    for pattern in FORGET_PATTERNS:
        match = re.search(pattern, msg)
        if match:
            content = match.group(1).strip() if match.lastindex else ""
            return ("forget", content)
    
    return (None, None)


async def handle_memory_command(command: str, content: str, conversation_id: str) -> str:
    now = datetime.now(timezone.utc).isoformat()

    if command == "remember":
        # Force ingest the specific content
        if not is_duplicate_memory(content, col_facts):
            col_facts.upsert(
                ids       = [f"explicit_fact_{uuid.uuid4().hex[:8]}"],
                documents = [content],
                metadatas = [{
                    "source":     "explicit",
                    "conv_id":    conversation_id,
                    "type":       "fact",
                    "created_at": now,
                }]
            )
            return f"Got it, I'll remember that: *{content}*"
        else:
            return f"I already have that stored."

    elif command == "forget":
        deleted = 0
        if content:
            # Search for the closest matching memory and delete it
            for col in [col_facts, col_knowledge, col_style]:
                try:
                    results = col.query(
                        query_texts=[content],
                        n_results=3,
                        include=["distances"]
                    )
                    if results["ids"] and results["ids"][0]:
                        for id_, dist in zip(results["ids"][0], results["distances"][0]):
                            if dist < 0.3:  # only delete if it's a close match
                                col.delete(ids=[id_])
                                deleted += 1
                except Exception:
                    pass
        
        if deleted:
            return f"Done, I've removed that from my memory."
        else:
            return f"I couldn't find anything close enough to that in my memory to remove."

# ── Document ingestion ────────────────────────────────────────────────────────
def ingest_document(doc_id: str, filename: str, text: str) -> int:
    chunks = chunk_text(text)
    now = datetime.now(timezone.utc).isoformat()

    for i, chunk in enumerate(chunks):
        col_documents.upsert(
            ids       = [f"doc_{doc_id}_chunk_{i}"],
            documents = [chunk],
            metadatas = [{
                "source":     "document",
                "doc_id":     doc_id,
                "filename":   filename,
                "chunk_idx":  str(i),
                "created_at": now,
            }]
        )

    return len(chunks)

def extract_text(file_bytes: bytes, filename: str) -> str:
    if filename.lower().endswith(".pdf"):
        reader = PdfReader(io.BytesIO(file_bytes))
        pages = [page.extract_text() or "" for page in reader.pages]
        return "\n\n".join(pages)
    else:
        return file_bytes.decode("utf-8", errors="ignore")

# ── Web search ────────────────────────────────────────────────────────────────
WEB_TRIGGERS = [
    r"\b(latest|recent|current news|today|tonight)\b",
    r"\b(news|weather|score|results?|standings?)\b",
    r"\b(search|look up|google)\b",
    r"\b(released?|launched?|announced?)\b",
    r"\b(final|winner|champion|trophy)\b",
]

def needs_web_search(message: str) -> bool:
    msg = message.lower()

    # Exclude messages that are clearly about internal data
    internal_signals = [
        "macro", "recipe", "ingredient", "meal", "my ", "i am", "i'm",
        "analyse deal", "analyze deal", "practice pitch", "remember",
        "what have i", "what did i", "shopify",
    ]
    if any(s in msg for s in internal_signals):
        return False

    return any(re.search(p, msg) for p in WEB_TRIGGERS)

async def web_search(query: str, max_results: int = 4) -> str:
    try:
        import asyncio
        results = await asyncio.to_thread(
            lambda: list(DDGS().text(query, max_results=max_results))
        )
        if not results:
            return ""
        lines = ["[Web search results:]"]
        for r in results:
            title = r.get("title", "")
            body  = r.get("body", "")
            url   = r.get("href", "")
            lines.append(f"• {title}: {body} ({url})")
        return "\n".join(lines)
    except Exception:
        return ""

# ── RAG: search memory ────────────────────────────────────────────────────────
def search_memory(query: str, n=5) -> str:
    """
    Pull summaries for relevant categories first, then fill with
    specific chunks from all collections. Summaries give the big
    picture; chunks give the detail.
    """
    candidates = []
    now_ts = datetime.now(timezone.utc).timestamp()

    # ── Step 1: Pull relevant summaries ──────────────────────────────────────
    try:
        summary_hits = col_summaries.query(
            query_texts=[query],
            n_results=3,
            include=["documents", "distances", "metadatas"]
        )
        for doc, dist, meta in zip(
            summary_hits["documents"][0],
            summary_hits["distances"][0],
            summary_hits["metadatas"][0]
        ):
            if not doc.strip():
                continue
            relevance = 1 / (1 + dist)
            # Only include summaries that are actually relevant
            if relevance > 0.3:
                candidates.append((
                    relevance * 1.4,  # boost summaries above raw chunks
                    f"Summary ({meta.get('category', 'general')})",
                    doc.strip()
                ))
    except Exception as e:
        print(f"Summary search error: {e}")

    # ── Step 2: Pull specific chunks from all collections ─────────────────────
    for col, label in [
        (col_knowledge, "Past conversation"),
        (col_facts,     "Things you've said"),
        (col_style,     "Your writing style"),
        (col_documents, "Your documents"),
    ]:
        try:
            hits = col.query(
                query_texts=[query],
                n_results=n,
                include=["documents", "distances", "metadatas"]
            )
            docs      = hits["documents"][0]
            distances = hits["distances"][0]
            metadatas = hits["metadatas"][0]

            for doc, dist, meta in zip(docs, distances, metadatas):
                if not doc.strip():
                    continue

                relevance = 1 / (1 + dist)

                try:
                    created_ts = datetime.fromisoformat(meta.get("created_at", "")).timestamp()
                    age_days   = (now_ts - created_ts) / 86400
                    recency    = math.exp(-age_days / 365)
                except Exception:
                    recency = 0.5

                type_boost = 1.15 if label in ("Things you've said", "Your documents") else 1.0
                final_score = (relevance * 0.7 + recency * 0.3) * type_boost
                candidates.append((final_score, label, doc.strip()))

        except Exception as e:
            print(f"search_memory error ({label}): {e}")

    # ── Step 3: Sort, filter, return ──────────────────────────────────────────
    candidates.sort(key=lambda x: x[0], reverse=True)
    top = [(label, doc) for score, label, doc in candidates if score > 0.15][:10]

    if not top:
        return ""

    return "\n\n".join(f"[{label}]: {doc}" for label, doc in top)

# ── System prompt ─────────────────────────────────────────────────────────────
def build_system_prompt(memory_context: str, web_context: str, sheets_context: str = "") -> str:
    base = """You are Lucchese, the personal AI of Alex Hammond.

Alex runs PTPreps — a meal prep business in the UK selling high protein meals in standard and bulking portions, available as one-time purchases or subscriptions via Shopify.
Alex is a personal trainer, into bodybuilding and martial arts, and is building this AI to automate his business and act as his most knowledgeable ally.

You know Alex well. Speak to him like a straight-talking, highly knowledgeable friend — not an assistant trying to please him.

Be direct and assertive. State things confidently without hedging.
Never use phrases like "it seems", "perhaps", "you might want to", "it could be", "I think", or "possibly" — if you know something, say it. If you don't, say so plainly.
Don't soften opinions or pad answers with disclaimers.
Don't be sycophantic — never open with praise or affirmations like "great question" or "absolutely".
When Alex is wrong or off track, say so directly and explain why. Challenge ideas that deserve to be challenged.
Match Alex's tone — casual, direct, no fluff.
Don't repeat yourself or over-explain.
Always end your response with a short, relevant question to keep the conversation moving.
Never guess or fabricate information about PTPreps, recipes, or macros — only use the Google Sheets data provided. If something isn't in the data, say so.
If you don't know something current like sports results, news, or prices — say so honestly.
When you use web search results, cite them naturally.
For ANY question about meals, ingredients, macros, or allergens — ONLY use the Google Sheets data provided. If a meal is not in the Sheets data, say "I don't have that meal in our current menu."
DOCUMENT GENERATION:
When the user asks you to write something as a document, Word doc, plan, programme, report,
or anything they'd want to save and use offline — generate the FULL content using proper
markdown structure. You MUST use markdown heading syntax:
  # Main Title
  ## Section Heading
  ### Subsection
  - bullet points for lists
  **bold** for key terms
  1. numbered steps where order matters

Then end your reply with exactly this marker on its own line:
[GENERATE_DOC: <short_descriptive_filename_no_extension>]
Example: [GENERATE_DOC: training_programme_week1]

IMPORTANT: Always use # and ## heading syntax. Never write section names as plain text.
Only use this marker when the content is genuinely document-worthy (structured plans,
programmes, checklists, reports). Not for short conversational answers."""


    sections = [base]

    if memory_context:
        sections.append(f"""IMPORTANT — MEMORY INSTRUCTIONS:
The following are real things Alex has said in past conversations. This IS your memory.
When Alex asks "what have I said about X" or "do I remember X", use ONLY what's below.
Do NOT say you don't have access to past conversations — you do, it's right here.
Speak about it naturally: "You've mentioned...", "You talked about...", "You said..."

--- Alex's Past Conversations & Facts ---
{memory_context}
---
Use this to inform your response where relevant, but don't quote it back verbatim.""")

    if sheets_context:
        sections.append(f"""Live data from PTPREPS Google Sheets:
---
{sheets_context}
---
Use this for any questions about recipes, ingredients, macros, or allergens.""")

    if web_context:
        sections.append(f"""Current information from the web:
---
{web_context}
---
Use ONLY this data for current events. Do not mix with your training knowledge.""")

    return "\n\n".join(sections)

# ── Request models ────────────────────────────────────────────────────────────
class ChatRequest(BaseModel):
    message:         str
    history:         Optional[list] = []
    conversation_id: Optional[str]  = None
    deep:            Optional[bool] = False

class ShopifyMealRequest(BaseModel):
    meal_name: str

# ── Shopify helpers ───────────────────────────────────────────────────────────
def calc_macros(meal):
    totals = {"kcal": 0, "protein": 0, "carbs": 0, "fat": 0}
    for i in meal['ingredients']:
        try:
            totals["kcal"]    += float(i['kcal'] or 0)
            totals["protein"] += float(i['protein'] or 0)
            totals["carbs"]   += float(i['carbs'] or 0)
            totals["fat"]     += float(i['fat'] or 0)
        except Exception:
            pass
    return totals

def normalise_meal_name(name: str) -> str:
    return name.lower().replace(' & ', ' and ').replace('&', 'and').strip()

async def _add_meal_to_shopify(meal_name: str) -> tuple:
    standard_recipes = get_recipes('Standard Recipe')
    bulking_recipes  = get_recipes('Bulking Recipes')
    matched = next((m for m in standard_recipes if normalise_meal_name(m) == normalise_meal_name(meal_name)), None)
    if not matched:
        return f"Couldn't find '{meal_name}' in the recipe sheet. Check the exact meal name and try again.", None
    standard_macros = calc_macros(standard_recipes[matched])
    bulking_macros  = calc_macros(bulking_recipes[matched]) if matched in bulking_recipes else standard_macros
    created = await _create_meal_products(matched, standard_macros, bulking_macros)
    return None, {"matched": matched, "standard_macros": standard_macros, "bulking_macros": bulking_macros, "created": created}

# ── Deal analyser ─────────────────────────────────────────────────────────────
MORTGAGE_RATE = 0.055  # 5.5% — update as rates change
STRESS_TEST_RATE = 0.08  # lenders stress test at ~8%

def analyse_deal(text: str) -> str:
    t = text.lower()
    
    # Find all £ amounts in the text
    # Handles: £200k, £200,000, £1,200, £950
    amounts = []
    for m in re.finditer(r'£([\d,]+)(k)?', t):
        val = int(m.group(1).replace(',', ''))
        if m.group(2):  # has 'k'
            val *= 1000
        amounts.append((m.start(), val))
    
    if not amounts:
        return "I need at least a purchase price and monthly rent. Try: 'analyse deal: 3 bed Basildon £200k asking, £1,200/month rent, single let'"
    
    # Price = largest amount
    price = max(amounts, key=lambda x: x[1])[1]
    
    # Rent = smallest amount (or amount near 'rent'/'month'/'pcm')
    rent_match = re.search(r'£([\d,]+)(k)?\s*(?:/month|pcm|per month|month|rent|/mo)', t)
    if rent_match:
        monthly_rent = int(rent_match.group(1).replace(',', ''))
        if rent_match.group(2):
            monthly_rent *= 1000
    else:
        # Fall back to smallest £ amount that isn't the price
        others = [v for _, v in amounts if v != price]
        monthly_rent = min(others) if others else None
    
    if not monthly_rent:
        return "I need an expected monthly rent to analyse this deal."
    
    # rest of the function unchanged from here...

    # Extract rooms (for HMO)
    rooms_match = re.search(r'(\d+)\s*(?:bed|room|bedroom)', t)
    rooms = int(rooms_match.group(1)) if rooms_match else None

    # Detect strategy
    is_hmo = 'hmo' in t or (rooms and rooms >= 4)
    is_r2r = 'r2r' in t or 'rent to rent' in t

    # For HMO, multiply per-room rent by number of rooms
    if is_hmo and rooms and monthly_rent < 2000:
        monthly_rent = monthly_rent * rooms

    if not price and not is_r2r:
        return "I need a purchase price to analyse this deal. Try: 'analyse deal: 3 bed Basildon £200k asking, £1,200/month rent, single let'"

    if not monthly_rent:
        return "I need an expected monthly rent to analyse this deal."

    lines = ["**Deal Analysis**\n"]

    if is_r2r:
        # R2R analysis — no purchase price needed
        landlord_rent = price or 0
        gross_monthly = monthly_rent
        expenses = gross_monthly * 0.15  # ~15% for bills/maintenance in R2R
        net_monthly = gross_monthly - landlord_rent - expenses
        lines += [
            f"**Strategy:** Rent to Rent",
            f"**Rent paid to landlord:** £{landlord_rent:,}/month",
            f"**Income from tenants:** £{gross_monthly:,}/month",
            f"**Estimated expenses (bills/maintenance):** £{expenses:,.0f}/month",
            f"**Net monthly cashflow:** £{net_monthly:,.0f}/month",
            f"**Annual profit:** £{net_monthly * 12:,.0f}",
            "",
            f"**Verdict:** {'✓ Positive cashflow — worth exploring' if net_monthly > 0 else '✗ Negative cashflow — numbers dont stack'}",
        ]
        return "\n".join(lines)

    annual_rent = monthly_rent * 12

    # Gross yield
    gross_yield = (annual_rent / price) * 100

    # Running costs
    if is_hmo:
        cost_pct = 0.35  # HMO has higher costs — bills, management, maintenance
    else:
        cost_pct = 0.25  # single let — agent fees, maintenance, insurance, voids

    annual_costs = annual_rent * cost_pct
    net_annual = annual_rent - annual_costs
    net_yield = (net_annual / price) * 100

    # Mortgage costs (75% LTV buy to let)
    ltv = 0.75
    loan = price * ltv
    deposit = price * 0.25
    monthly_mortgage = (loan * MORTGAGE_RATE) / 12  # interest only
    stress_mortgage = (loan * STRESS_TEST_RATE) / 12

    # Cashflow
    monthly_costs = annual_costs / 12
    monthly_cashflow = monthly_rent - monthly_mortgage - monthly_costs

    # SDLT (investment property — includes 3% surcharge)
    def calc_sdlt(p):
        if p <= 125000:
            sdlt = p * 0.03
        elif p <= 250000:
            sdlt = 125000 * 0.03 + (p - 125000) * 0.05
        elif p <= 925000:
            sdlt = 125000 * 0.03 + 125000 * 0.05 + (p - 250000) * 0.08
        else:
            sdlt = 125000 * 0.03 + 125000 * 0.05 + 675000 * 0.08 + (p - 925000) * 0.13
        return sdlt

    sdlt = calc_sdlt(price)

    # Total cash required
    total_cash = deposit + sdlt + 2000  # rough legal/valuation fees

    # Stress test — does rent cover stressed mortgage?
    stress_coverage = monthly_rent / stress_mortgage
    passes_stress = stress_coverage >= 1.25  # most lenders require 125% coverage

    # Equity release scenario
    equity_needed = deposit
    equity_property_value = equity_needed / 0.25  # what property value needed to release this as 25% equity
    remortgage_amount = equity_needed

    lines += [
        f"**Strategy:** {'HMO' if is_hmo else 'Single Let'}",
        f"**Purchase price:** £{price:,}",
        f"**Monthly rent:** £{monthly_rent:,}",
        "",
        "**Yield**",
        f"Gross yield: {gross_yield:.1f}%",
        f"Net yield (after {int(cost_pct*100)}% costs): {net_yield:.1f}%",
        f"{'✓ Strong yield' if gross_yield >= 8 else '✓ Decent yield' if gross_yield >= 6 else '⚠ Low yield — check numbers'} for {'HMO' if is_hmo else 'single let'}",
        "",
        "**Mortgage (75% LTV, interest only)**",
        f"Loan amount: £{loan:,.0f}",
        f"Deposit required: £{deposit:,.0f}",
        f"Monthly mortgage (~{MORTGAGE_RATE*100:.1f}%): £{monthly_mortgage:,.0f}",
        f"Stress test rate ({STRESS_TEST_RATE*100:.0f}%): £{stress_mortgage:,.0f}/month",
        f"Stress test coverage: {stress_coverage:.2f}x {'✓ Passes' if passes_stress else '✗ Fails — lender may decline'}",
        "",
        "**Monthly Cashflow**",
        f"Rent: £{monthly_rent:,}",
        f"Mortgage: -£{monthly_mortgage:,.0f}",
        f"Running costs: -£{monthly_costs:,.0f}",
        f"Net cashflow: £{monthly_cashflow:,.0f}/month {'✓' if monthly_cashflow > 0 else '✗'}",
        f"Annual cashflow: £{monthly_cashflow * 12:,.0f}",
        "",
        "**Upfront Costs**",
        f"Deposit (25%): £{deposit:,.0f}",
        f"SDLT (inc. 3% surcharge): £{sdlt:,.0f}",
        f"Legal/valuation (est.): £2,000",
        f"Total cash required: £{total_cash:,.0f}",
        "",
        "**Equity Release Angle**",
        f"To fund this deposit via remortgage, the existing property needs",
        f"at least £{equity_needed:,.0f} available equity (at 75% LTV).",
        "",
        "**Verdict**",
    ]

    # Overall verdict
    issues = []
    positives = []

    if gross_yield >= 8:
        positives.append("strong gross yield")
    elif gross_yield >= 6:
        positives.append("decent gross yield")
    else:
        issues.append(f"low gross yield of {gross_yield:.1f}%")

    if monthly_cashflow > 200:
        positives.append("solid monthly cashflow")
    elif monthly_cashflow > 0:
        positives.append("positive but thin cashflow")
    else:
        issues.append("negative cashflow")

    if not passes_stress:
        issues.append("fails lender stress test — income may need to be higher or deposit larger")

    if issues:
        lines.append(f"⚠ Concerns: {', '.join(issues)}")
    if positives:
        lines.append(f"✓ Positives: {', '.join(positives)}")

    if monthly_cashflow > 0 and passes_stress and gross_yield >= 6:
        lines.append("\n**Overall: Worth pursuing — book a viewing and get broker involved.**")
    elif monthly_cashflow > 0 and gross_yield >= 5:
        lines.append("\n**Overall: Marginal — negotiate the price down or find ways to increase rent.**")
    else:
        lines.append("\n**Overall: Hard to make work at this price — move on or negotiate hard.**")

    return "\n".join(lines)

    # ── Property pitch role play ──────────────────────────────────────────────────
def get_roleplay_session(conversation_id: str) -> dict | None:
    con = get_con()
    row = con.execute(
        "SELECT exchanges FROM roleplay_sessions WHERE conversation_id = ?",
        (conversation_id,)
    ).fetchone()
    con.close()
    return {"exchanges": row["exchanges"]} if row else None

def upsert_roleplay_session(conversation_id: str, exchanges: int):
    now = datetime.now(timezone.utc).isoformat()
    con = get_con()
    con.execute("""
        INSERT INTO roleplay_sessions (conversation_id, exchanges, created_at)
        VALUES (?, ?, ?)
        ON CONFLICT(conversation_id) DO UPDATE SET exchanges = excluded.exchanges
    """, (conversation_id, exchanges, now))
    con.commit()
    con.close()

def delete_roleplay_session(conversation_id: str) -> dict:
    session = get_roleplay_session(conversation_id)
    con = get_con()
    con.execute("DELETE FROM roleplay_sessions WHERE conversation_id = ?", (conversation_id,))
    con.commit()
    con.close()
    return session or {}

ROLEPLAY_SYSTEM = """You are playing the role of a family member called Carol — Alex's aunt. You are 58 years old, live in a 3 bed semi in Essex that you own outright, bought it around 18 years ago. You are warm but sharp and slightly dismissive of Alex — not cruel, just the way older family members can be with the younger ones who've always done things a bit differently.

Background:
- Alex and his brother have always been seen as the black sheep of the family — they do things their own way, which the family respects but doesn't always take seriously
- You don't really follow what Alex does for work — you know he's self employed and does something with food and fitness but you don't ask much about it
- There have been some deaths in the family recently so gatherings are less frequent — Alex comes to see you one on one which you appreciate
- You have a financial advisor called Michael who all the family use — he keeps things simple and safe, pensions, savings, that sort of thing. When Alex's mum mentioned property investment to Michael he wasn't keen and that's stuck with you
- Your plan for the house is simple — stay in it, then leave it to the kids in the will. That's what you've always assumed you'd do and it feels safe and right
- You are not financially literate in terms of investing or money working for you — nobody ever taught you that and you've never needed to think about it. But you would never admit that — you have a big ego and would rather seem skeptical than confused
- You love Alex but see him as young and a bit of a dreamer — one of his schemes that probably won't go anywhere

Your personality:
- Sharp, warm but dismissive — "that sounds complicated love", "are you sure you know what you're talking about?", "I've worked too hard for this house", "Michael looked into this sort of thing and said it wasn't worth the hassle"
- If Alex mentions leaving more to the kids or the will, your ears prick up slightly — that's your soft spot
- If Alex is vague or gets defensive, get more skeptical
- If Alex is calm, specific, and respectful of your concerns, warm up slightly — but don't cave easily
- Never write actions, stage directions, or descriptions. No "Carol sips her tea" or "she narrows her eyes". Plain dialogue only.
- Never write Alex's lines or describe what Alex does. Only respond as Carol.
- Keep responses short — 2-4 sentences max. Real conversations don't have speeches.
- Go off topic occasionally — mention another family member, something happening locally, but nothing invented about Alex's personal life

Start the conversation as Carol, reacting to Alex having just said he wanted to talk to her about something regarding her property."""

async def run_roleplay(conversation_id: str, user_msg: str, history: list) -> str:
    msg = user_msg.lower().strip()

    # End session
    if any(x in msg for x in ["end practice", "end role play", "end roleplay", "stop practice"]):
        session = delete_roleplay_session(conversation_id)
        exchanges = session.get("exchanges", 0)

        feedback_prompt = f"""You were just playing Carol in a property pitch role play with Alex. Step out of character completely and give honest, constructive feedback on how Alex handled the conversation.

The conversation had {exchanges} exchanges.

Conversation history:
{chr(10).join([f"{m['role'].upper()}: {m['content']}" for m in history[-20:]])}

Give feedback on:
1. Opening — did he open naturally or did it feel like a sales pitch?
2. Clarity — did he explain things simply without jargon?
3. Confidence — did he sound like he knew what he was talking about?
4. Handling objections — how did he deal with skepticism or difficult questions?
5. Listening — did he respond to what Carol actually said or just talk at her?
6. Next steps — did he leave with a clear next step or did it fizzle out?

Be direct and specific. Point to actual moments in the conversation. End with one thing he should focus on for next time."""

        async with httpx.AsyncClient(timeout=120) as client:
            res = await client.post(OLLAMA_URL, json={
                "model": MODEL_FAST,
                "messages": [{"role": "user", "content": feedback_prompt}],
                "stream": False,
            })
            return "**[End of practice — here's your feedback:]**\n\n" + res.json()["message"]["content"]

    # Track session — increment exchange count
    session = get_roleplay_session(conversation_id)
    exchanges = (session["exchanges"] if session else 0) + 1
    upsert_roleplay_session(conversation_id, exchanges)

    # Build messages for Carol
    messages = [{"role": "system", "content": ROLEPLAY_SYSTEM}]
    for m in history:
        messages.append({"role": m["role"], "content": m["content"]})
    messages.append({"role": "user", "content": user_msg})

    async with httpx.AsyncClient(timeout=120) as client:
        res = await client.post(OLLAMA_URL, json={
            "model": MODEL_FAST,
            "messages": messages,
            "stream": False,
        })
        return res.json()["message"]["content"]

# ── Chat endpoint ─────────────────────────────────────────────────────────────
@app.post("/chat")
async def chat(req: ChatRequest):
    conversation_id = req.conversation_id or str(uuid.uuid4())

    def stream_plain_reply(reply: str, web_search_used: bool = False, auto_ingested: bool = False):
        """Wrap a plain string reply in the NDJSON streaming format the frontend expects."""
        async def generator():
            yield json.dumps({"type": "meta", "conversation_id": conversation_id, "web_search_used": web_search_used}) + "\n"
            yield json.dumps({"type": "token", "content": reply}) + "\n"
            yield json.dumps({"type": "done", "auto_ingested": auto_ingested}) + "\n"
        return StreamingResponse(generator(), media_type="application/x-ndjson")

    # ── Shopify command intercept ─────────────────────────────────────────────
    shopify_match = re.search(r'shopify add (.+)|add (.+) to shopify', req.message.lower())
    if shopify_match:
        meal_name = (shopify_match.group(1) or shopify_match.group(2)).strip()
        error, result = await _add_meal_to_shopify(meal_name)
        if error:
            reply = error
        else:
            reply = f"Done! Created 4 products for {result['matched']} on Shopify:\n"
            for p in result['created']: reply += f"  ✓ {p['title']}\n"

    # ── Memory command intercept ──────────────────────────────────────────────
    command, content = detect_memory_command(req.message)
    if command:
        reply = await handle_memory_command(command, content, conversation_id)
        save_message(conversation_id, "user", req.message)
        save_message(conversation_id, "assistant", reply)
        return stream_plain_reply(reply, auto_ingested=command == "remember")

    # ── Deal analysis intercept ───────────────────────────────────────────────
    if req.message.lower().startswith("analyse deal:") or req.message.lower().startswith("analyze deal:"):
        reply = analyse_deal(req.message)
        save_message(conversation_id, "user", req.message)
        save_message(conversation_id, "assistant", reply)
        return stream_plain_reply(reply)

    # ── Role play intercept ───────────────────────────────────────────────────
    msg_lower = req.message.lower().strip()
    is_active_roleplay = get_roleplay_session(conversation_id) is not None
    starts_roleplay = any(x in msg_lower for x in ["practice pitch", "role play property", "roleplay property", "start practice"])

    if starts_roleplay or is_active_roleplay:
        reply = await run_roleplay(conversation_id, req.message, req.history)
        # Mark session as started if not already
        if starts_roleplay and not is_active_roleplay:
            upsert_roleplay_session(conversation_id, 0)
        save_message(conversation_id, "user", req.message)
        save_message(conversation_id, "assistant", reply)
        return stream_plain_reply(reply)

    # ── Normal chat flow ──────────────────────────────────────────────────────
    did_search = needs_web_search(req.message)
    web = await web_search(req.message) if did_search else ""
    memory     = "" if did_search else search_memory(req.message)
    sheets     = get_menu_context(req.message)

    messages = [{"role": "system", "content": build_system_prompt(memory, web, sheets)}]
    messages += req.history
    messages.append({"role": "user", "content": req.message})

    save_message(conversation_id, "user", req.message)

    async def stream_response():
        full_reply = []
        auto_ingest = False
 
        yield json.dumps({
            "type": "meta",
            "conversation_id": conversation_id,
            "web_search_used": did_search,
        }) + "\n"
 
        try:
            if CHAT_PROVIDER == "claude":
                system_prompt = messages[0]["content"] if messages[0]["role"] == "system" else ""
                chat_messages = [m for m in messages if m["role"] != "system"]
                async with httpx.AsyncClient(timeout=300) as client:
                    res = await client.post(
                        "https://api.anthropic.com/v1/messages",
                        headers={
                            "x-api-key": ANTHROPIC_API_KEY,
                            "anthropic-version": "2023-06-01",
                            "content-type": "application/json",
                        },
                        json={
                            "model": "claude-sonnet-4-6",
                            "max_tokens": 4096,
                            "system": system_prompt,
                            "messages": chat_messages,
                        }
                    )
                if res.status_code != 200:
                    raise Exception(f"Anthropic API error {res.status_code}: {res.text[:200]}")
                reply_text = res.json()["content"][0]["text"]
                full_reply.append(reply_text)
                yield json.dumps({"type": "token", "content": reply_text}) + "\n"
 
            else:
                # ── Ollama (existing behaviour, unchanged) ────────────────────
                async with httpx.AsyncClient(timeout=300) as client:
                    async with client.stream("POST", OLLAMA_URL, json={
                        "model":    MODEL_DEEP if req.deep else MODEL_FAST,
                        "messages": messages,
                        "stream":   True,
                    }) as response:
                        async for line in response.aiter_lines():
                            if not line.strip():
                                continue
                            try:
                                chunk = json.loads(line)
                                token = chunk.get("message", {}).get("content", "")
                                if token:
                                    full_reply.append(token)
                                    yield json.dumps({"type": "token", "content": token}) + "\n"
                                if chunk.get("done"):
                                    break
                            except Exception:
                                continue
 
        except Exception as e:
            print(f"stream_response error ({CHAT_PROVIDER}): {e}")
            yield json.dumps({"type": "token", "content": "\n\n[Response error — please try again]"}) + "\n"
 
        # ── Post-processing (identical regardless of provider) ────────────────
        reply = "".join(full_reply)
        if reply:
            save_message(conversation_id, "assistant", reply)
 
            user_corrections = ['we already', 'actually', "that's wrong", 'not quite',
                                'to clarify', "we don't", 'we do', 'i am', "i'm not"]
            force_ingest = any(s in req.message.lower() for s in user_corrections)
 
            if force_ingest:
                auto_ingest = True
                await ingest_exchange(conversation_id, req.message, reply)
            elif not did_search:
                auto_ingest = should_ingest(req.message, reply)
                print(f"Should ingest: {auto_ingest} | Message: '{req.message[:40]}'")
                if auto_ingest:
                    await ingest_exchange(conversation_id, req.message, reply)
 
        yield json.dumps({"type": "done", "auto_ingested": auto_ingest}) + "\n"
    return StreamingResponse(stream_response(), media_type="application/x-ndjson")


# ── Cross reference endpoint ──────────────────────────────────────────────
@app.get("/cross-reference")
def cross_reference(sheets: str = None, threshold: float = 0.8):
    """
    Cross-reference raw ingredients against recipe ingredients.
    sheets: comma-separated sheet names (e.g., "Standard Recipe,Bulking Recipes"), defaults to both
    threshold: similarity threshold 0-1 for fuzzy matching (default 0.8)
    """
    try:
        if sheets:
            sheets_list = [s.strip() for s in sheets.split(',')]
        else:
            sheets_list = None

        result = cross_reference_ingredients(sheets_list, threshold)
        return result
    except Exception as e:
        return {"error": str(e)}, 500


# ── Shopify endpoint ──────────────────────────────────────────────────────────
@app.post("/shopify/add-meal")
async def add_meal_to_shopify(req: ShopifyMealRequest):
    error, result = await _add_meal_to_shopify(req.meal_name.strip())
    if error:
        return {"error": error}
    return {
        "success": True,
        "meal":             result["matched"],
        "standard_macros":  result["standard_macros"],
        "bulking_macros":   result["bulking_macros"],
        "products_created": result["created"],
    }

# ── Upload endpoint ───────────────────────────────────────────────────────────
@app.post("/upload")
async def upload_document(file: UploadFile = File(...)):
    contents = await file.read()
    filename = file.filename or "unknown"

    text = extract_text(contents, filename)
    if not text.strip():
        return {"error": "Could not extract text from file"}, 400

    doc_id      = uuid.uuid4().hex
    chunk_count = ingest_document(doc_id, filename, text)

    now = datetime.now(timezone.utc).isoformat()
    file_type = "pdf" if filename.lower().endswith(".pdf") else "text"
    con = get_con()
    con.execute(
        "INSERT INTO documents (id, filename, file_type, chunk_count, created_at) VALUES (?, ?, ?, ?, ?)",
        (doc_id, filename, file_type, chunk_count, now)
    )
    con.commit()
    con.close()

    return {
        "doc_id":      doc_id,
        "filename":    filename,
        "chunk_count": chunk_count,
        "message":     f"Ingested {chunk_count} chunks from {filename}",
    }

@app.get("/documents")
def list_documents():
    con = get_con()
    rows = con.execute(
        "SELECT id, filename, file_type, chunk_count, created_at FROM documents ORDER BY created_at DESC"
    ).fetchall()
    con.close()
    return [dict(r) for r in rows]

@app.delete("/documents/{doc_id}")
def delete_document(doc_id: str):
    results = col_documents.get(where={"doc_id": doc_id})
    if results["ids"]:
        col_documents.delete(ids=results["ids"])

    con = get_con()
    con.execute("DELETE FROM documents WHERE id = ?", (doc_id,))
    con.commit()
    con.close()

    return {"deleted": doc_id}

# ── History endpoints ─────────────────────────────────────────────────────────
@app.get("/conversations")
def list_conversations():
    con = get_con()
    rows = con.execute(
        "SELECT id, title, created_at, updated_at FROM conversations ORDER BY updated_at DESC"
    ).fetchall()
    con.close()
    return [dict(r) for r in rows]

@app.get("/conversations/{conversation_id}")
def get_conversation(conversation_id: str):
    con = get_con()
    rows = con.execute(
        "SELECT role, content, created_at FROM messages WHERE conversation_id = ? ORDER BY created_at",
        (conversation_id,)
    ).fetchall()
    con.close()
    return [dict(r) for r in rows]

@app.delete("/conversations/{conversation_id}")
def delete_conversation(conversation_id: str):
    # Remove from ChromaDB memory
    for col in [col_knowledge, col_facts, col_style]:
        try:
            results = col.get(where={"conv_id": conversation_id})
            if results["ids"]:
                col.delete(ids=results["ids"])
        except Exception as e:
            print(f"delete_conversation chroma error: {e}")


    # Remove from SQLite
    con = get_con()
    con.execute("DELETE FROM messages WHERE conversation_id = ?", (conversation_id,))
    con.execute("DELETE FROM conversations WHERE id = ?", (conversation_id,))
    con.commit()
    con.close()
    return {"deleted": conversation_id}

# ── Feedback endpoint ─────────────────────────────────────────────────────────
class FeedbackRequest(BaseModel):
    conversation_id: str
    user_message: str
    assistant_reply: str
    rating: str  # "good" or "bad"

@app.post("/feedback")
async def feedback(req: FeedbackRequest):
    if req.rating == "good":
        await ingest_exchange(req.conversation_id, req.user_message, req.assistant_reply)
        return {"ingested": True}
    else:
        # On bad rating, try to remove if already ingested
        try:
            results = col_knowledge.get(where={"conv_id": req.conversation_id})
            if results["ids"]:
                col_knowledge.delete(ids=results["ids"])
        except Exception as e:
            print(f"feedback delete error: {e}")
        return {"ingested": False}

DOC_MARKER_RE = re.compile(r"\[GENERATE_DOC:\s*([^\]]+)\]", re.IGNORECASE)

@app.post("/generate-doc")
async def generate_doc(req: dict):
    """
    Called by the frontend when it detects a [GENERATE_DOC: ...] marker.
    Body: { "content": "...", "title": "..." }
    Returns: { "token": "...", "filename": "..." }
    """
    content = req.get("content", "").strip()
    title   = req.get("title", "document").strip()

    if not content:
        return {"error": "No content provided"}

    try:
        filepath = markdown_to_docx(content, title)
        token    = uuid.uuid4().hex
        filename = Path(filepath).name

        with _doc_lock:
            _doc_store[token] = filepath

        # Auto-cleanup after 15 minutes
        async def cleanup():
            import asyncio
            await asyncio.sleep(900)
            with _doc_lock:
                path = _doc_store.pop(token, None)
            if path and os.path.exists(path):
                os.remove(path)

        import asyncio
        asyncio.create_task(cleanup())

        return {"token": token, "filename": filename}

    except Exception as e:
        print(f"generate_doc error: {e}")
        return {"error": str(e)}


@app.get("/download/{token}")
async def download_doc(token: str):
    with _doc_lock:
        filepath = _doc_store.get(token)

    if not filepath or not os.path.exists(filepath):
        return JSONResponse({"error": "Document not found or expired"}, status_code=404)

    filename = Path(filepath).name
    return FileResponse(
        path        = filepath,
        filename    = filename,
        media_type  = "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


# ── Shared audio transcription helper ────────────────────────────────────────
async def transcribe_audio(file: UploadFile) -> str:
    contents = await file.read()
    content_type = file.content_type or ""
    if "mp4" in content_type or "m4a" in content_type: suffix = ".mp4"
    elif "ogg" in content_type:                         suffix = ".ogg"
    elif "wav" in content_type:                         suffix = ".wav"
    else:                                               suffix = ".webm"

    tmp_fd, tmp_path = tempfile.mkstemp(suffix=suffix)
    try:
        with os.fdopen(tmp_fd, "wb") as tmp:
            tmp.write(contents)
        result = whisper_model.transcribe(tmp_path)
    finally:
        if os.path.exists(tmp_path):
            os.remove(tmp_path)
    return result["text"].strip()

@app.post("/transcribe")
async def transcribe(file: UploadFile = File(...)):
    try:
        text = await transcribe_audio(file)
        return {"text": text}
    except Exception as e:
        print(f"Transcribe error: {e}")
        return {"error": str(e)}


# ── Text to Speech endpoint ───────────────────────────────────────────────────
class TTSRequest(BaseModel):
    text: str

@app.post("/tts")
async def text_to_speech(req: TTSRequest):
    try:
        audio = el_client.text_to_speech.convert(
            voice_id=ELEVENLABS_VOICE_ID,
            text=req.text,
            model_id="eleven_turbo_v2_5",
            output_format="mp3_44100_128",
        )
        audio_bytes = b"".join(audio)
        return Response(
            content=audio_bytes,
            media_type="audio/mpeg",
            headers={"Content-Disposition": "inline; filename=speech.mp3"}
        )
    except Exception as e:
        return {"error": str(e)}

@app.post("/voice-chat")
async def voice_chat(file: UploadFile = File(...), conversation_id: str = None):
    try:
        # 1. Transcribe
        user_text = await transcribe_audio(file)
        if not user_text:
            return JSONResponse(content={"error": "no speech detected"})

        conv_id = conversation_id or str(uuid.uuid4())

        # 2. Build context and get reply
        memory = search_memory(user_text)
        system = build_system_prompt(memory, "", "")

        # Load conversation history for context
        history = []
        try:
            con = get_con()
            rows = con.execute(
                "SELECT role, content FROM messages WHERE conversation_id = ? ORDER BY created_at DESC LIMIT 20",
                (conv_id,)
            ).fetchall()
            con.close()
            history = [{"role": r["role"], "content": r["content"]} for r in reversed(rows)]
        except Exception as e:
            print(f"voice-chat history load error: {e}")

        messages = [{"role": "system", "content": system}]
        messages += history
        messages.append({"role": "user", "content": user_text})

        async with httpx.AsyncClient(timeout=300) as client:
            res = await client.post(OLLAMA_URL, json={
                "model": MODEL_FAST,
                "messages": messages,
                "stream": False,
            })
        reply_text = res.json()["message"]["content"]

        # 3. Save + ingest
        save_message(conv_id, "user",      user_text)
        save_message(conv_id, "assistant", reply_text)
        if should_ingest(user_text, reply_text):
            await ingest_exchange(conv_id, user_text, reply_text)

        # 4. TTS — strip emojis, collect audio, base64 encode
        reply_clean = re.sub(r'[^\x00-\x7F\u00C0-\u024F\u1E00-\u1EFF]+', ' ', reply_text).strip()
        sentences   = [s.strip() for s in re.split(r'(?<=[.!?])\s+', reply_clean) if s.strip()]

        merged, buf = [], ""
        for s in sentences:
            buf = (buf + " " + s).strip()
            if len(buf) > 80:
                merged.append(buf)
                buf = ""
        if buf:
            merged.append(buf)

        audio_bytes = b""
        for chunk_text in merged:
            try:
                for chunk in el_client.text_to_speech.convert(
                    voice_id=ELEVENLABS_VOICE_ID,
                    text=chunk_text,
                    model_id="eleven_turbo_v2_5",
                    output_format="mp3_44100_128",
                ):
                    audio_bytes += chunk
            except Exception as e:
                print(f"TTS error: {e}")

        return JSONResponse(
            content={
                "transcript": user_text,
                "reply":      reply_text,
                "conv_id":    conv_id,
                "audio_b64":  base64.b64encode(audio_bytes).decode("utf-8"),
            },
            media_type="application/json; charset=utf-8"
        )

    except Exception as e:
        print(f"Voice chat error: {e}")
        return JSONResponse(content={"error": str(e)})

@app.get("/debug/memory")
def debug_memory():
    results = col_knowledge.get(include=["documents", "metadatas"])
    items = []
    for doc, meta in zip(results["documents"], results["metadatas"]):
        items.append({
            "text": doc[:100],
            "created_at": meta.get("created_at"),
            "source": meta.get("source"),
            "type": meta.get("type"),
        })
    items.sort(key=lambda x: x["created_at"] or "", reverse=True)
    return items

# ── Admin endpoints ───────────────────────────────────────────────────────────
@app.get("/admin/stats")
def admin_stats():
    now_ts = datetime.now(timezone.utc).timestamp()
    stats = {}
    for col, name in [(col_knowledge, "knowledge"), (col_facts, "facts"), (col_style, "style")]:
        results = col.get(include=["metadatas"])
        metas = results["metadatas"]
        source_counts = {}
        category_counts = {}
        for m in metas:
            src = m.get("source", "unknown")
            cat = m.get("category", "general")
            source_counts[src] = source_counts.get(src, 0) + 1
            category_counts[cat] = category_counts.get(cat, 0) + 1
        stats[name] = {
            "total": len(metas),
            "by_source": source_counts,
            "by_category": category_counts,
        }
    return stats

@app.get("/admin/recent")
def admin_recent(limit: int = 20):
    results = col_facts.get(include=["documents", "metadatas"])
    items = []
    for doc, meta in zip(results["documents"], results["metadatas"]):
        items.append({
            "text": doc[:120],
            "source": meta.get("source", "unknown"),
            "category": meta.get("category", "general"),
            "created_at": meta.get("created_at", ""),
            "conv_id": meta.get("conv_id", ""),
        })
    items.sort(key=lambda x: x["created_at"] or "", reverse=True)
    return items[:limit]

@app.get("/admin/search")
def admin_search(q: str, n: int = 10):
    results = col_facts.query(
        query_texts=[q],
        n_results=n,
        include=["documents", "metadatas", "distances"]
    )
    items = []
    for doc, meta, dist in zip(results["documents"][0], results["metadatas"][0], results["distances"][0]):
        items.append({
            "text": doc[:200],
            "source": meta.get("source", "unknown"),
            "category": meta.get("category", "general"),
            "created_at": meta.get("created_at", ""),
            "relevance": round(1 / (1 + dist), 3),
        })
    return items

@app.delete("/admin/memory")
def admin_delete_memory(source: str):
    deleted = 0
    for col in [col_knowledge, col_facts, col_style]:
        try:
            results = col.get(where={"source": source})
            if results["ids"]:
                col.delete(ids=results["ids"])
                deleted += len(results["ids"])
        except Exception as e:
            print(f"admin_delete_memory error: {e}")
    return {"deleted": deleted, "source": source}

# ── Summarise endpoint ────────────────────────────────────────────────────────
@app.post("/admin/summarise")
async def admin_summarise():
    now = datetime.now(timezone.utc).isoformat()
    results_log = {}

    # Pull all facts and knowledge grouped by category
    all_docs = []
    for col, label in [(col_facts, "facts"), (col_knowledge, "knowledge")]:
        try:
            results = col.get(include=["documents", "metadatas"])
            for doc, meta in zip(results["documents"], results["metadatas"]):
                if doc.strip():
                    all_docs.append({
                        "text":     doc.strip(),
                        "category": meta.get("category", "general"),
                        "label":    label,
                    })
        except Exception as e:
            import traceback
            print(f"Summarise error ({category}): {e}")
            print(traceback.format_exc())
            results_log[category] = f"error: {e}"

    # Group by category
    from collections import defaultdict
    by_category = defaultdict(list)
    for d in all_docs:
        by_category[d["category"]].append(d["text"])

    # Summarise each category with enough entries
    MIN_ENTRIES = 5
    for category, texts in by_category.items():
        if len(texts) < MIN_ENTRIES:
            results_log[category] = f"skipped ({len(texts)} entries, need {MIN_ENTRIES})"
            continue

        # Take the most recent/relevant sample — cap at 60 chunks to avoid huge prompts
        sample = texts[:15]
        combined = "\n\n---\n\n".join(sample)

        prompt = f"""You are summarising Alex Hammond's memory about the topic: "{category}".

The following are real things Alex has said or discussed in past conversations.
Write a dense, coherent summary of everything you know about Alex in relation to this topic.
Include specific details, preferences, history, goals, and any recurring themes.
Write it as a factual third-person profile, like a knowledgeable assistant briefing someone about Alex.
Do not include generic filler — only specific, useful information.
Keep it under 400 words.

--- Memory entries ---
{combined}
---

Summary:"""

        try:
            async with httpx.AsyncClient(timeout=120) as client:
                res = await client.post(OLLAMA_URL, json={
                    "model": MODEL_FAST,
                    "messages": [{"role": "user", "content": prompt}],
                    "stream": False,
                })
            summary_text = res.json()["message"]["content"].strip()

            # Upsert into summaries collection — one entry per category
            col_summaries.upsert(
                ids       = [f"summary_{category}"],
                documents = [summary_text],
                metadatas = [{
                    "category":     category,
                    "entry_count":  str(len(texts)),
                    "last_updated": now,
                    "type":         "summary",
                }]
            )
            results_log[category] = f"ok ({len(texts)} entries summarised)"
            print(f"Summarised: {category} ({len(texts)} entries)")

        except Exception as e:
            print(f"Summarise error ({category}): {e}")
            results_log[category] = f"error: {e}"

    return {"summarised": results_log, "total_categories": len(results_log)}

@app.get("/admin/summaries")
def admin_summaries():
    try:
        results = col_summaries.get(include=["documents", "metadatas"])
        items = []
        for doc, meta in zip(results["documents"], results["metadatas"]):
            items.append({
                "category":     meta.get("category"),
                "entry_count":  meta.get("entry_count"),
                "last_updated": meta.get("last_updated"),
                "summary":      doc,
            })
        items.sort(key=lambda x: int(x["entry_count"] or 0), reverse=True)
        return items
    except Exception as e:
        return {"error": str(e)}


# ── Health check ──────────────────────────────────────────────────────────────
@app.get("/health")
def health():
    return {
        "status":    "ok",
        "knowledge": col_knowledge.count(),
        "facts":     col_facts.count(),
        "style":     col_style.count(),
        "documents": col_documents.count(),
    }