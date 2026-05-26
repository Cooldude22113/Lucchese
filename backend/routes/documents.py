"""
documents.py
Word document generation for Lucchese.
Handles: markdown to docx conversion, token-based download store, cleanup.
"""

import re
import uuid
import os
import asyncio
import threading
from datetime import datetime, timezone
from pathlib import Path
from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ── Storage ───────────────────────────────────────────────────────────────────
DOCS_DIR = Path("./generated_docs")
DOCS_DIR.mkdir(exist_ok=True)

_doc_store: dict[str, str] = {}
_doc_lock = threading.Lock()

# ── Inline formatting ─────────────────────────────────────────────────────────
def _add_inline_formatting(para, text: str):
    """Handle **bold** and *italic* inline markdown within a paragraph."""
    pattern = re.compile(r"(\*\*.*?\*\*|\*.*?\*)")
    parts   = pattern.split(text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run      = para.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run        = para.add_run(part[1:-1])
            run.italic = True
        else:
            para.add_run(part)


def _looks_like_section_header(line: str) -> bool:
    """
    Detect plain-text section headers like "Property Basics:" or "Exit Strategy:"
    when the model doesn't use markdown heading syntax.
    Rules: ends with colon, 1-5 words, no sentence punctuation, starts with capital.
    """
    s = line.strip()
    if not s.endswith(":"):
        return False
    if len(s) > 50:
        return False
    if not s[0].isupper():
        return False
    word_count = len(s.rstrip(":").split())
    if word_count > 5:
        return False
    if re.search(r"[,\.!?']", s.rstrip(":")):
        return False
    return True


# ── Markdown to docx ──────────────────────────────────────────────────────────
def markdown_to_docx(content: str, title: str) -> str:
    """
    Convert markdown-ish text to a .docx file.
    Handles: # ## ### headings, plain "Section:" headers, - bullets,
    1. numbered lists, **bold**, *italic*, --- rules.
    Returns the filepath of the saved document.
    """
    doc = DocxDocument()

    # ── Page margins ──────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.2)
        section.right_margin  = Inches(1.2)

    # ── Title ─────────────────────────────────────────────────────────────────
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title_para.add_run(title.replace("_", " ").title())
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

    date_para = doc.add_paragraph()
    date_run  = date_para.add_run(
        datetime.now(timezone.utc).strftime("Generated %d %B %Y")
    )
    date_run.font.size = Pt(9)
    date_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    doc.add_paragraph()  # spacer

    # ── Body ──────────────────────────────────────────────────────────────────
    for line in content.split("\n"):
        s = line.strip()

        if not s:
            doc.add_paragraph()
            continue

        # H1: # Heading
        if s.startswith("# "):
            p   = doc.add_paragraph()
            run = p.add_run(s[2:].strip())
            run.bold = True
            run.font.size = Pt(15)
            run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after  = Pt(4)

        # H2: ## Heading
        elif s.startswith("## "):
            p   = doc.add_paragraph()
            run = p.add_run(s[3:].strip())
            run.bold = True
            run.font.size = Pt(13)
            run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after  = Pt(3)

        # H3: ### Heading
        elif s.startswith("### "):
            p   = doc.add_paragraph()
            run = p.add_run(s[4:].strip())
            run.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x55)
            p.paragraph_format.space_before = Pt(8)

        # Plain-text section header fallback: "Section Name:"
        elif _looks_like_section_header(s):
            p   = doc.add_paragraph()
            run = p.add_run(s.rstrip(":"))
            run.bold = True
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after  = Pt(3)

        # Bullet: - item or * item
        elif s.startswith(("- ", "* ")):
            p = doc.add_paragraph(style="List Bullet")
            _add_inline_formatting(p, s[2:].strip())

        # Numbered: 1. item
        elif re.match(r"^\d+\.\s", s):
            p = doc.add_paragraph(style="List Number")
            _add_inline_formatting(p, re.sub(r"^\d+\.\s*", "", s))

        # Horizontal rule
        elif s in ("---", "***", "___"):
            p   = doc.add_paragraph()
            run = p.add_run("─" * 60)
            run.font.color.rgb = RGBColor(0xcc, 0xcc, 0xcc)
            run.font.size = Pt(8)

        # Normal paragraph
        else:
            p = doc.add_paragraph()
            _add_inline_formatting(p, s)
            p.paragraph_format.space_after = Pt(4)

    # ── Save ──────────────────────────────────────────────────────────────────
    safe_name = re.sub(r"[^\w\-]", "_", title)[:60]
    filename  = f"{safe_name}_{uuid.uuid4().hex[:6]}.docx"
    filepath  = DOCS_DIR / filename
    doc.save(str(filepath))
    return str(filepath)


# ── Token store helpers ───────────────────────────────────────────────────────
def store_doc_token(filepath: str) -> tuple[str, str]:
    """Store a filepath under a token for download. Returns (token, filename)."""
    token    = uuid.uuid4().hex
    filename = Path(filepath).name
    with _doc_lock:
        _doc_store[token] = filepath
    return token, filename


def get_doc_path(token: str) -> str | None:
    """Retrieve the filepath for a token. Returns None if not found."""
    with _doc_lock:
        return _doc_store.get(token)


async def schedule_cleanup(token: str, delay_seconds: int = 900):
    """Auto-delete a generated doc after delay_seconds (default 15 min)."""
    await asyncio.sleep(delay_seconds)
    with _doc_lock:
        path = _doc_store.pop(token, None)
    if path and os.path.exists(path):
        os.remove(path)
